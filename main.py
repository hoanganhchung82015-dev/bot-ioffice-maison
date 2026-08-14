import os
import re
import logging
import asyncio
from typing import List, Dict, Any
from datetime import datetime
from io import BytesIO

from aiohttp import web
from playwright.async_api import async_playwright
import google.generativeai as genai
from telegram import Update
from telegram.ext import Application, CommandHandler, ContextTypes
from pypdf import PdfReader

# Thư viện tạo file Word
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ==========================================
# 1. CẤU HÌNH LOGGING & BIẾN MÔI TRƯỜNG
# ==========================================
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)

def sanitize_url(raw: str) -> str:
    if not raw:
        return "https://thptmaison.vnptioffice.vn"
    clean = re.sub(r'[^\x20-\x7E]', '', raw).strip().strip('"').strip("'")
    if not clean.startswith(("http://", "https://")):
        clean = f"https://{clean}"
    return clean.rstrip('/')

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()
IOFFICE_URL = sanitize_url(os.getenv("IOFFICE_URL", "https://thptmaison.vnptioffice.vn"))
IOFFICE_USERNAME = os.getenv("IOFFICE_USERNAME", "").strip()
IOFFICE_PASSWORD = os.getenv("IOFFICE_PASSWORD", "").strip()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
PORT = int(os.getenv("PORT", 10000))

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# ==========================================
# 2. XỬ LÝ ĐỌC FILE PDF
# ==========================================
def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Trích xuất toàn bộ văn bản từ file PDF."""
    try:
        reader = PdfReader(BytesIO(pdf_bytes))
        full_text = []
        max_pages = min(len(reader.pages), 12)
        for i in range(max_pages):
            text = reader.pages[i].extract_text()
            if text:
                full_text.append(text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Lỗi đọc PDF: {e}")
        return ""

# ==========================================
# 3. TẠO FILE WORD 7 CỘT CHI TIẾT
# ==========================================
def set_cell_background(cell, fill_hex):
    """Tô màu nền cho ô trong bảng Word."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_docx_report(data_list: List[Dict[str, Any]]) -> BytesIO:
    """Tạo file Word 7 cột chi tiết báo cáo BGH."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BẢNG TỔNG HỢP CHI TIẾT PHÂN CÔNG XỬ LÝ VĂN BẢN ĐẾN")
    title_run.font.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Times New Roman'
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Trường THPT Mai Sơn — Ngày xuất báo cáo: {datetime.now().strftime('%d/%m/%Y')}\n")
    sub_run.font.italic = True
    sub_run.font.size = Pt(10.5)
    sub_run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = [
        ("STT", Inches(0.4)),
        ("Số / Ký hiệu & Ngày VB", Inches(1.1)),
        ("Tên văn bản / Trích yếu", Inches(1.8)),
        ("Người chỉ đạo", Inches(1.1)),
        ("Đơn vị / Người thực hiện", Inches(1.1)),
        ("Nội dung chỉ đạo & Yêu cầu sản phẩm", Inches(1.8)),
        ("Thời hạn hoàn thành", Inches(0.9))
    ]

    hdr_cells = table.rows[0].cells
    for idx, (text, width) in enumerate(headers):
        hdr_cells[idx].width = width
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[idx], "003366")

    for item in data_list:
        row_cells = table.add_row().cells
        row_data = [
            (str(item.get("stt", "")), WD_ALIGN_PARAGRAPH.CENTER, True),
            (item.get("so_hieu", "Đang cập nhật"), WD_ALIGN_PARAGRAPH.CENTER, True),
            (item.get("title", ""), WD_ALIGN_PARAGRAPH.LEFT, False),
            (item.get("chi_dao", ""), WD_ALIGN_PARAGRAPH.LEFT, True),
            (item.get("thuc_hien", ""), WD_ALIGN_PARAGRAPH.LEFT, False),
            (item.get("ket_qua", ""), WD_ALIGN_PARAGRAPH.LEFT, False),
            (item.get("han_chot", ""), WD_ALIGN_PARAGRAPH.CENTER, False)
        ]

        for idx, (text, align, is_bold) in enumerate(row_data):
            row_cells[idx].width = headers[idx][1]
            p = row_cells[idx].paragraphs[0]
            p.alignment = align
            run = p.add_run(text)
            run.font.size = Pt(9.0)
            run.font.name = 'Times New Roman'
            run.font.bold = is_bold
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# ==========================================
# 4. CÀO DỮ LIỆU IOFFICE & PDF
# ==========================================
async def scan_and_process_ioffice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    status_msg = await update.message.reply_text("🔍 **Đang khởi động kết nối VNPT iOffice...**", parse_mode="Markdown")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=["--no-sandbox", "--disable-setuid-sandbox", "--disable-dev-shm-usage", "--disable-gpu"]
        )
        context_browser = await browser.new_context(viewport={"width": 1366, "height": 768})
        page = await context_browser.new_page()

        try:
            # 1. ĐĂNG NHẬP
            await page.goto(IOFFICE_URL, timeout=60000, wait_until="domcontentloaded")
            username_sel = "input[name='username'], input[id='username'], input[type='text']"
            password_sel = "input[name='password'], input[id='password'], input[type='password']"
            submit_sel = "button[type='submit'], input[type='submit'], .btn-login, #btnLogin"

            if await page.is_visible(username_sel):
                await page.fill(username_sel, IOFFICE_USERNAME)
                await page.fill(password_sel, IOFFICE_PASSWORD)
                await page.click(submit_sel)
                await page.wait_for_load_state("networkidle")

            # 2. VÀO MỤC DUYỆT VĂN BẢN ĐẾN
            await status_msg.edit_text("📂 **Đang mở danh sách Duyệt văn bản đến...**", parse_mode="Markdown")
            vb_den_menu = page.locator("a:has-text('Văn bản đến'), span:has-text('Văn bản đến')").first
            if await vb_den_menu.is_visible():
                await vb_den_menu.click()
                await page.wait_for_timeout(1000)

            duyet_vb_menu = page.locator("a:has-text('Duyệt văn bản đến'), span:has-text('Duyệt văn bản đến')").first
            if await duyet_vb_menu.is_visible():
                await duyet_vb_menu.click()
                await page.wait_for_load_state("networkidle")

            await page.wait_for_timeout(3000)

            # 3. QUÉT TOÀN BỘ VĂN BẢN TRONG BẢNG
            rows = await page.query_selector_all("table tbody tr")
            doc_items = []

            for r in rows:
                link_elem = await r.query_selector("td a, a.doc-title, td:nth-child(4) a, td:nth-child(3) a")
                if link_elem:
                    title_text = (await link_elem.inner_text()).strip().replace("\n", " ")
                    if len(title_text) > 8 and not any(k in title_text.lower() for k in ["trích yếu", "toggle navigation", "vb đến"]):
                        doc_items.append({"title": title_text})

            total_docs = len(doc_items)
            if total_docs == 0:
                await status_msg.edit_text("ℹ️ **Không tìm thấy văn bản nào cần duyệt trong danh sách.**", parse_mode="Markdown")
                await browser.close()
                return

            await status_msg.edit_text(f"✅ Phát hiện **{total_docs}** văn bản đến. Bắt đầu đọc chi tiết PDF & phân công...", parse_mode="Markdown")

            parsed_results = []

            # 4. DUYỆT LẦN LƯỢT TẤT CẢ VĂN BẢN
            for idx in range(total_docs):
                try:
                    trich_yeu = doc_items[idx]["title"]
                    await status_msg.edit_text(
                        f"⏳ **[{idx+1}/{total_docs}]** Đang đọc file PDF văn bản:\n📄 _{trich_yeu[:70]}..._",
                        parse_mode="Markdown"
                    )

                    # Re-query danh sách để lấy element chính xác
                    current_rows = await page.query_selector_all("table tbody tr")
                    target_link = None
                    for cr in current_rows:
                        txt = (await cr.inner_text()).strip()
                        if trich_yeu[:20] in txt:
                            target_link = await cr.query_selector("td a, a.doc-title, td:nth-child(4) a, td:nth-child(3) a")
                            break

                    pdf_text_content = ""

                    if target_link:
                        try:
                            # Thử click download
                            async with page.expect_download(timeout=4000) as download_info:
                                await target_link.click()
                            download = await download_info.value
                            path = await download.path()
                            with open(path, "rb") as f:
                                pdf_bytes = f.read()
                            pdf_text_content = extract_text_from_pdf_bytes(pdf_bytes)
                        except Exception:
                            # Nếu mở trang chi tiết
                            await page.wait_for_timeout(2000)
                            pdf_btn = await page.query_selector("a[href*='.pdf'], a:has-text('.pdf'), .file-attach a")
                            if pdf_btn:
                                try:
                                    async with page.expect_download(timeout=4000) as download_info2:
                                        await pdf_btn.click()
                                    download2 = await download_info2.value
                                    path2 = await download2.path()
                                    with open(path2, "rb") as f:
                                        pdf_bytes2 = f.read()
                                    pdf_text_content = extract_text_from_pdf_bytes(pdf_bytes2)
                                except Exception:
                                    pass

                            if not pdf_text_content:
                                body_elem = await page.query_selector(".doc-detail-content, .panel-body, #content-detail, body")
                                if body_elem:
                                    pdf_text_content = await body_elem.inner_text()

                            # Quay lại bảng
                            back_btn = page.locator("button:has-text('Quay lại'), a:has-text('Quay lại')").first
                            if await back_btn.is_visible():
                                await back_btn.click()
                            else:
                                await page.go_back()
                            await page.wait_for_timeout(1500)

                    # Phân tích sâu bằng Gemini 1.5 Flash
                    ai_dict = await analyze_document_with_gemini(trich_yeu, pdf_text_content)
                    ai_dict["stt"] = len(parsed_results) + 1
                    ai_dict["title"] = trich_yeu
                    parsed_results.append(ai_dict)

                except Exception as row_err:
                    logger.error(f"Lỗi xử lý văn bản {idx+1}: {row_err}")
                    continue

            # 5. XUẤT VÀ GỬI FILE WORD
            if not parsed_results:
                await status_msg.edit_text("❌ **Không trích xuất được dữ liệu văn bản nào.**", parse_mode="Markdown")
                await browser.close()
                return

            await status_msg.edit_text("📝 **Đang đóng gói file Word báo cáo 7 cột chi tiết...**", parse_mode="Markdown")
            
            docx_file = create_docx_report(parsed_results)
            filename = f"Bao_Cao_Chi_Tiet_iOffice_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

            await context.bot.send_document(
                chat_id=update.effective_chat.id,
                document=docx_file,
                filename=filename,
                caption=f"🎉 **Đã hoàn thành!** Đã bóc tách số hiệu, đọc PDF và lập báo cáo phân công chi tiết cho toàn bộ **{len(parsed_results)}** văn bản đến."
            )
            await status_msg.delete()

        except Exception as e:
            logger.error(f"Lỗi quy trình: {e}")
            await status_msg.edit_text(f"❌ **Lỗi trong quá trình cào dữ liệu**: `{e}`", parse_mode="Markdown")
        finally:
            await browser.close()

# ==========================================
# 5. GEMINI PHÂN TÍCH SÂU TỪNG FILE PDF
# ==========================================
async def analyze_document_with_gemini(title: str, full_pdf_text: str) -> Dict[str, str]:
    """Phân tích PDF để bóc tách Số hiệu, Người chỉ đạo, Đơn vị thực hiện, Yêu cầu và Hạn chót."""
    default_res = {
        "so_hieu": "Đang cập nhật",
        "chi_dao": "Hiệu trưởng Hoàng Anh Chung",
        "thuc_hien": "Các bộ phận liên quan",
        "han_chot": "Theo quy định",
        "ket_qua": "Thực hiện theo chỉ đạo tại văn bản"
    }

    if not GEMINI_API_KEY:
        return default_res

    clean_text = full_pdf_text[:4000] if full_pdf_text else "Không lấy được nội dung PDF, phân tích theo trích yếu."

    prompt = f"""
    Bạn là Thư ký Ban Giám hiệu Trường THPT Mai Sơn.
    Hãy đọc thật kỹ văn bản/file PDF sau để bóc tách thông tin lập bảng phân công xử lý văn bản đến:

    📌 **Trích yếu văn bản**: "{title}"
    📄 **Nội dung PDF chi tiết**:
    "{clean_text}"

    ================ QUY TẮC BÓC TÁCH THÔNG TIN ================
    1. **SỐ HIỆU & NGÀY VB**: Tìm số công văn/kế hoạch/quyết định và ngày ban hành (Ví dụ: "Số 125/KH-SGDĐT ngày 10/08/2026" hoặc "Số 45/QĐ-UBND"). Nếu không tìm thấy trong PDF thì ghi "Theo trích yếu".
    
    2. **NGƯỜI CHỈ ĐẠO**:
       - "Hiệu trưởng Hoàng Anh Chung": Công tác Đảng, Tổ chức cán bộ, Tài chính, Quy hoạch, Tài sản công/Đất đai, Chỉ đạo chung.
       - "PHT Lại Thế Dũng": Chuyên môn dạy học, Thi HSG/GVG, GDPT 2018, Chuyển đổi số, CNTT, Văn hóa/văn nghệ/thể thao học sinh.
       - "PHT CSVC & Lao động": Quản lý trang thiết bị, Cơ sở vật chất, Lao động, An ninh trật tự, PCCC, Y tế trường học.

    3. **ĐƠN VỊ / NGƯỜI THỰC HIỆN**: Ghi rõ các Tổ chuyên môn (Tổ Toán, Tổ Ngữ Văn...), Đoàn TN, Kế toán, Văn thư, GVCN, Tổ Trợ lý...

    4. **NỘI DUNG CHỈ ĐẠO & YÊU CẦU SẢN PHẨM**: Tóm tắt ngắn gọn 2-3 câu các nhiệm vụ cụ thể cần làm và sản phẩm đầu ra (Ví dụ: "Lập kế hoạch triển khai cuộc thi; Lập danh sách giáo viên tham gia tập huấn; Nộp báo cáo thống kê hiện trạng...").

    5. **THỜI HẠN HOÀN THÀNH**: Tìm đúng mốc thời gian chót phải hoàn thành/nộp báo cáo ghi trong PDF (Ví dụ: "Trước 16h00 ngày 25/08/2026"). Nếu không ghi rõ ngày thì ghi "Theo quy định".

    TRẢ VỀ DUY NHẤT 5 DÒNG THEO CÚ PHÁP CHÍNH XÁC:
    SỐ HIỆU: <Ghi số ký hiệu và ngày VB>
    CHỈ ĐẠO: <Ghi tên Lãnh đạo chỉ đạo>
    THỰC HIỆN: <Ghi đơn vị/người thực hiện>
    YÊU CẦU: <Ghi tóm tắt nội dung chỉ đạo & sản phẩm cần nộp>
    HẠN CHÓT: <Ghi thời hạn hoàn thành>
    """

    try:
        # Sử dụng model gemini-1.5-flash chuẩn định danh
        model = genai.GenerativeModel('gemini-1.5-flash')
        res = await asyncio.to_thread(model.generate_content, prompt)
        text = res.text if res.text else ""

        so_hieu = re.search(r"SỐ HIỆU:\s*(.*)", text)
        chi_dao = re.search(r"CHỈ ĐẠO:\s*(.*)", text)
        thuc_hien = re.search(r"THỰC HIỆN:\s*(.*)", text)
        yeu_cau = re.search(r"YÊU CẦU:\s*(.*)", text)
        han_chot = re.search(r"HẠN CHÓT:\s*(.*)", text)

        return {
            "so_hieu": so_hieu.group(1).strip() if so_hieu else default_res["so_hieu"],
            "chi_dao": chi_dao.group(1).strip() if chi_dao else default_res["chi_dao"],
            "thuc_hien": thuc_hien.group(1).strip() if thuc_hien else default_res["thuc_hien"],
            "ket_qua": yeu_cau.group(1).strip() if yeu_cau else default_res["ket_qua"],
            "han_chot": han_chot.group(1).strip() if han_chot else default_res["han_chot"],
        }
    except Exception as e:
        logger.error(f"Lỗi Gemini: {e}")
        return default_res

# ==========================================
# 6. KHỞI CHẠY BOT
# ==========================================
async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("👋 **Xin chào BGH THPT Mai Sơn!**\n\nGõ lệnh **/scan** để Bot đọc chi tiết từng file PDF trên iOffice và xuất file Word báo cáo phân công 7 cột.", parse_mode="Markdown")

async def scan_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await scan_and_process_ioffice(update, context)

async def handle_ping(request):
    return web.Response(text="Bot Running OK!")

async def main():
    if not TELEGRAM_BOT_TOKEN:
        logger.error("Thiếu TELEGRAM_BOT_TOKEN!")
        return

    app = web.Application()
    app.router.add_get('/', handle_ping)
    runner = web.AppRunner(app)
    await runner.setup()
    await web.TCPSite(runner, '0.0.0.0', PORT).start()

    telegram_app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    telegram_app.add_handler(CommandHandler("start", start_command))
    telegram_app.add_handler(CommandHandler("scan", scan_command))

    await telegram_app.initialize()
    await telegram_app.start()
    await telegram_app.updater.start_polling(drop_pending_updates=True)
    
    logger.info("Bot iOffice 7 cột đã sẵn sàng!")
    await asyncio.Event().wait()

if __name__ == "__main__":
    asyncio.run(main())
